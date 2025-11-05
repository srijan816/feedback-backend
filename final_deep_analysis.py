#!/usr/bin/env python3
"""
Ultra-deep analysis of ONE complete feedback file to understand ALL structure details
"""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import json

def analyze_complete_document(file_path):
    """Analyze every detail of a docx file"""
    print(f"=" * 100)
    print(f"ULTRA-DEEP ANALYSIS OF: {file_path}")
    print(f"=" * 100)

    doc = Document(file_path)

    # Analyze paragraphs
    print("\n" + "-" * 100)
    print("PARAGRAPHS")
    print("-" * 100)
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            print(f"\nParagraph {i}:")
            print(f"  Text: {para.text[:200]}")
            print(f"  Style: {para.style.name}")

    # Analyze tables in detail
    print("\n" + "-" * 100)
    print("TABLES (DETAILED)")
    print("-" * 100)

    for t_idx, table in enumerate(doc.tables, 1):
        print(f"\n{'=' * 80}")
        print(f"TABLE {t_idx}")
        print(f"{'=' * 80}")
        print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")

        for r_idx, row in enumerate(table.rows):
            print(f"\n  Row {r_idx}:")
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    # Show first 150 characters
                    preview = cell_text if len(cell_text) <= 150 else cell_text[:150] + "..."
                    print(f"    Cell [{r_idx},{c_idx}]: {preview}")

def main():
    # Analyze a few diverse files
    files_to_analyze = [
        "/home/ubuntu/apps/feedback-backend/Mai/G7-12-PSD-III-Wednesdays-6pm/Alex Zhu/Feedback - 2.3 - Alex Zhu.docx",
        "/home/ubuntu/apps/feedback-backend/Mai/G7-12-PSD-II- Tuesdays-6pm/Emma Lai/Unit 1/Feedback - 1.2 - Emma Lai.docx"
    ]

    for file_path in files_to_analyze:
        try:
            analyze_complete_document(file_path)
            print("\n\n")
        except Exception as e:
            print(f"ERROR analyzing {file_path}: {e}")

if __name__ == '__main__':
    main()
