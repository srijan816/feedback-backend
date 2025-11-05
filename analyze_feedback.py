#!/usr/bin/env python3
"""
Analyze feedback docx files to understand rubrics and teacher comment patterns
"""

import os
from docx import Document
import json
from pathlib import Path

def extract_text_from_docx(file_path):
    """Extract all text from a docx file"""
    try:
        doc = Document(file_path)
        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Extract tables
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables_data.append(table_data)

        return {
            'paragraphs': full_text,
            'tables': tables_data,
            'full_text': '\n'.join(full_text)
        }
    except Exception as e:
        return {'error': str(e)}

def analyze_folder_structure(base_path):
    """Analyze the folder structure"""
    structure = {
        'classes': [],
        'total_students': 0,
        'total_feedback_files': 0
    }

    base = Path(base_path)

    for class_dir in base.iterdir():
        if class_dir.is_dir() and not class_dir.name.startswith('.'):
            class_info = {
                'name': class_dir.name,
                'students': []
            }

            for student_dir in class_dir.iterdir():
                if student_dir.is_dir() and not student_dir.name.startswith('.'):
                    # Count feedback files for this student
                    feedback_files = list(student_dir.rglob('*.docx'))
                    class_info['students'].append({
                        'name': student_dir.name,
                        'feedback_count': len(feedback_files),
                        'files': [str(f.relative_to(base)) for f in feedback_files]
                    })
                    structure['total_students'] += 1
                    structure['total_feedback_files'] += len(feedback_files)

            structure['classes'].append(class_info)

    return structure

def main():
    mai_folder = '/home/ubuntu/apps/feedback-backend/Mai'

    print("=" * 80)
    print("ANALYZING MAI FOLDER STRUCTURE")
    print("=" * 80)

    structure = analyze_folder_structure(mai_folder)

    print(f"\nTotal Classes: {len(structure['classes'])}")
    print(f"Total Students: {structure['total_students']}")
    print(f"Total Feedback Files: {structure['total_feedback_files']}")

    for class_info in structure['classes']:
        print(f"\n\nClass: {class_info['name']}")
        print(f"  Students: {len(class_info['students'])}")

        # Show sample students
        for student in class_info['students'][:3]:
            print(f"    - {student['name']}: {student['feedback_count']} feedback files")

    print("\n" + "=" * 80)
    print("ANALYZING SAMPLE FEEDBACK FILES")
    print("=" * 80)

    # Get a diverse sample of feedback files
    all_files = []
    for class_info in structure['classes']:
        for student in class_info['students']:
            all_files.extend(student['files'])

    # Sample from different students and units
    sample_files = []
    seen_patterns = set()

    for file_path in all_files:
        # Extract pattern (e.g., "1.1", "1.2", "2.3")
        pattern = file_path.split(' - ')[-2] if ' - ' in file_path else ''

        if pattern and pattern not in seen_patterns:
            sample_files.append(file_path)
            seen_patterns.add(pattern)

        if len(sample_files) >= 15:  # Analyze 15 diverse samples
            break

    # Analyze each sample file
    for i, rel_path in enumerate(sample_files, 1):
        file_path = os.path.join(mai_folder, rel_path)
        print(f"\n{'=' * 80}")
        print(f"FILE {i}/{len(sample_files)}: {rel_path}")
        print(f"{'=' * 80}")

        content = extract_text_from_docx(file_path)

        if 'error' in content:
            print(f"ERROR: {content['error']}")
            continue

        # Print paragraphs
        if content['paragraphs']:
            print("\n--- PARAGRAPHS ---")
            for j, para in enumerate(content['paragraphs'][:20], 1):  # First 20 paragraphs
                print(f"{j}. {para}")

        # Print tables
        if content['tables']:
            print("\n--- TABLES ---")
            for t_idx, table in enumerate(content['tables'], 1):
                print(f"\nTable {t_idx}:")
                for row in table:
                    print(f"  | {' | '.join(row)}")

        print("\n" + "-" * 80)

if __name__ == '__main__':
    main()
